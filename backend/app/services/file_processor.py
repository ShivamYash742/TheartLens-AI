"""services/file_processor.py — Text extraction from PDF, DOCX, TXT, CSV files."""
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

def extract_text(file_path, mime_type: str) -> str:
    path = Path(file_path)
    logger.info("Extracting text from %s (mime=%s)", path.name, mime_type)
    try:
        if mime_type == "application/pdf":
            return _extract_pdf(path)
        elif "wordprocessingml" in mime_type or path.suffix.lower() == ".docx":
            return _extract_docx(path)
        elif mime_type == "text/plain" or path.suffix.lower() == ".txt":
            return _extract_txt(path)
        elif mime_type in ("text/csv", "application/csv", "text/x-csv") or path.suffix.lower() == ".csv":
            return _extract_csv(path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
    except Exception as exc:
        logger.error("Extraction failed: %s", exc)
        raise RuntimeError(f"Text extraction failed: {exc}") from exc

def _extract_pdf(path: Path) -> str:
    import fitz
    doc = fitz.open(str(path))
    text = "\n".join(page.get_text() for page in doc).strip()
    doc.close()
    if not text:
        logger.warning("PDF %s returned empty — may need OCR", path.name)
    return text

def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)

def _extract_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")

def _extract_csv(path: Path) -> str:
    import pandas as pd
    df = pd.read_csv(path, dtype=str, keep_default_na=False)
    rows = []
    for _, row in df.iterrows():
        rows.append("; ".join(f"{c}: {v}" for c, v in row.items() if v))
    return "\n".join(rows)
